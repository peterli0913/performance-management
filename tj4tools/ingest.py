"""把上传的文件 / 文件夹 / zip 统一读成"表"和"文档"。

设计目标：
  * 一次上传里混着 xlsx / xls / csv / docx / pdf / zip 都能吃下；
  * 提取失败必须显式记录（文件名 + 原因），绝不静默跳过；
  * 同一份文件重复上传按 sha256 去重。
"""

from __future__ import annotations

import csv
import datetime as _dt
import hashlib
import io
import os
import posixpath
import re
import zipfile
from dataclasses import dataclass, field

from .normalize import clean_text

EXCEL_SUFFIXES = (".xlsx", ".xlsm", ".xltx", ".xltm")
LEGACY_EXCEL_SUFFIXES = (".xls",)
CSV_SUFFIXES = (".csv", ".tsv", ".txt")
WORD_SUFFIXES = (".docx",)
PDF_SUFFIXES = (".pdf",)
ZIP_SUFFIXES = (".zip",)

MAX_ZIP_ENTRIES = 2000
MAX_MEMBER_BYTES = 200 * 1024 * 1024


@dataclass
class SheetTable:
    """一张二维表。``header`` 已定位到真正的表头行。"""

    file_name: str
    sheet_name: str
    header: list[str]
    rows: list[list[object]]
    header_row_index: int = 1

    @property
    def label(self) -> str:
        return f"{self.file_name} / {self.sheet_name}"

    @property
    def n_rows(self) -> int:
        return len(self.rows)


@dataclass
class TextDocument:
    file_name: str
    kind: str
    text: str


@dataclass
class IngestResult:
    tables: list[SheetTable] = field(default_factory=list)
    documents: list[TextDocument] = field(default_factory=list)
    files: list[dict] = field(default_factory=list)
    problems: list[str] = field(default_factory=list)
    raw_files: dict[str, bytes] = field(default_factory=dict)

    def add_problem(self, message: str) -> None:
        self.problems.append(message)


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _zip_member_name(member: zipfile.ZipInfo) -> str:
    """zip 里的中文名常是 GBK 编码，flag_bits 未置 0x800 时需要还原。"""
    name = member.filename
    if member.flag_bits & 0x800:
        return name
    try:
        return name.encode("cp437").decode("gbk")
    except (UnicodeDecodeError, UnicodeEncodeError):
        return name


def _safe_member_path(name: str) -> str | None:
    """拒绝目录穿越与绝对路径。"""
    normalized = posixpath.normpath(name.replace("\\", "/"))
    if normalized.startswith(("/", "../")) or normalized == ".." or os.path.isabs(normalized):
        return None
    return normalized


def ingest_files(items: list[tuple[str, bytes]]) -> IngestResult:
    """``items`` 为 (文件名, 字节) 列表；zip 会被递归展开。"""
    result = IngestResult()
    seen: set[str] = set()
    queue = list(items)
    while queue:
        name, data = queue.pop(0)
        digest = _sha256(data)
        if digest in seen:
            result.add_problem(f"跳过重复文件：{name}")
            continue
        seen.add(digest)
        suffix = os.path.splitext(name)[1].lower()
        if suffix in ZIP_SUFFIXES:
            queue.extend(_expand_zip(name, data, result))
            continue
        entry = {
            "file_name": name,
            "sha256": digest,
            "size": len(data),
            "kind": suffix.lstrip(".") or "unknown",
            "n_tables": 0,
        }
        try:
            if suffix in EXCEL_SUFFIXES:
                tables = read_excel(name, data)
            elif suffix in LEGACY_EXCEL_SUFFIXES:
                tables = read_legacy_excel(name, data)
            elif suffix in CSV_SUFFIXES:
                tables = read_csv(name, data)
            elif suffix in WORD_SUFFIXES:
                tables, document = read_docx(name, data)
                if document is not None:
                    result.documents.append(document)
            elif suffix in PDF_SUFFIXES:
                tables, document = read_pdf(name, data)
                if document is not None:
                    result.documents.append(document)
            else:
                result.add_problem(f"暂不支持的文件类型，已跳过：{name}")
                tables = []
        except Exception as exc:  # noqa: BLE001 - 逐文件容错，避免一份坏文件毁掉整批
            result.add_problem(f"解析失败：{name}（{type(exc).__name__}: {exc}）")
            tables = []
        entry["n_tables"] = len(tables)
        result.tables.extend(tables)
        result.files.append(entry)
        result.raw_files[name] = data
    return result


def _expand_zip(name: str, data: bytes, result: IngestResult) -> list[tuple[str, bytes]]:
    out: list[tuple[str, bytes]] = []
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile:
        result.add_problem(f"压缩包无法打开：{name}")
        return out
    with archive:
        members = archive.infolist()
        if len(members) > MAX_ZIP_ENTRIES:
            result.add_problem(f"压缩包条目过多（{len(members)}），只处理前 {MAX_ZIP_ENTRIES} 个：{name}")
            members = members[:MAX_ZIP_ENTRIES]
        for member in members:
            if member.is_dir():
                continue
            member_name = _zip_member_name(member)
            safe = _safe_member_path(member_name)
            if safe is None:
                result.add_problem(f"压缩包内路径非法，已跳过：{member_name}")
                continue
            base = posixpath.basename(safe)
            if base.startswith(".") or safe.startswith("__MACOSX"):
                continue
            if member.file_size > MAX_MEMBER_BYTES:
                result.add_problem(f"压缩包内文件过大，已跳过：{safe}")
                continue
            out.append((f"{name}::{safe}", archive.read(member)))
    return out


# --------------------------------------------------------------------------- #
# 各类型读取
# --------------------------------------------------------------------------- #


_HEADER_DATE_RE = re.compile(r"^\d{4}[-/.]\d{1,2}([-/.]\d{1,2})?$")


def _looks_like_label(value) -> bool:
    """表头单元格应该是文字标签，而不是日期或纯数字。"""
    if isinstance(value, (int, float, _dt.date, _dt.datetime)) and not isinstance(value, bool):
        return False
    text = clean_text(value)
    if not text:
        return False
    if _HEADER_DATE_RE.match(text):
        return False
    return not text.replace(".", "", 1).replace("-", "", 1).isdigit()


def pick_header_row(grid: list[list[object]], max_scan: int = 8) -> int:
    """在前若干行里挑最像表头的一行。

    "非空最多"是不够的：一线人员第 2 行有 31 个日期单元格，会盖过第 1 行的真表头。
    所以只给文字标签计正分，给日期/数字计负分，并轻微偏向靠前的行。
    """
    best_index, best_score = 0, float("-inf")
    for index, row in enumerate(grid[:max_scan]):
        labels = sum(1 for cell in row if _looks_like_label(cell))
        non_labels = sum(1 for cell in row if cell is not None and not _looks_like_label(cell) and clean_text(cell))
        if labels == 0 and non_labels == 0:
            continue
        score = labels - non_labels * 0.5 - index * 0.3
        if score > best_score:
            best_index, best_score = index, score
    return best_index


def _dedupe_header(header: list[str]) -> list[str]:
    out: list[str] = []
    seen: dict[str, int] = {}
    for index, raw in enumerate(header):
        name = clean_text(raw) or f"列{index + 1}"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 1
        out.append(name)
    return out


def read_excel(file_name: str, data: bytes) -> list[SheetTable]:
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=False)
    tables: list[SheetTable] = []
    try:
        for sheet in workbook.worksheets:
            grid = [list(row) for row in sheet.iter_rows(values_only=True)]
            if not grid:
                continue
            # 先在未填充的网格上定位表头，再决定填充范围：
            # 表头自身的纵向合并（如 A1:A2）不能向下填充，否则会多出一行"假数据"
            header_row = pick_header_row(grid) + 1
            header_span_end = header_row
            for merged in sheet.merged_cells.ranges:
                if merged.min_row == header_row and merged.max_row > header_span_end:
                    header_span_end = merged.max_row
            merged_fill(sheet, grid, skip_until_row=header_span_end)
            header = _dedupe_header(grid[header_row - 1])
            rows = [
                row
                for row in grid[header_span_end:]
                if any(cell is not None and clean_text(cell) for cell in row)
            ]
            tables.append(
                SheetTable(
                    file_name=file_name,
                    sheet_name=sheet.title,
                    header=header,
                    rows=rows,
                    header_row_index=header_row,
                )
            )
    finally:
        workbook.close()
    return tables


def merged_fill(sheet, grid: list[list[object]], skip_until_row: int = 0) -> None:
    """把纵向合并单元格的值向下填充，否则分组列会整片丢失。

    ``skip_until_row`` 之前起始的合并区不处理（用于跳过表头自身的合并）。
    """
    for merged in sheet.merged_cells.ranges:
        if merged.min_row == merged.max_row or merged.min_row <= skip_until_row:
            continue
        row_index = merged.min_row - 1
        col_index = merged.min_col - 1
        if row_index >= len(grid) or col_index >= len(grid[row_index]):
            continue
        value = grid[row_index][col_index]
        if value is None:
            continue
        for target in range(merged.min_row, merged.max_row):
            if target < len(grid) and col_index < len(grid[target]):
                grid[target][col_index] = value


def read_legacy_excel(file_name: str, data: bytes) -> list[SheetTable]:
    try:
        import xlrd  # type: ignore
    except ImportError as exc:  # pragma: no cover - 依赖可选
        raise RuntimeError("读取 .xls 需要安装 xlrd<2.0，请改存为 .xlsx") from exc
    book = xlrd.open_workbook(file_contents=data)
    tables = []
    for sheet in book.sheets():
        grid = [[sheet.cell_value(r, c) for c in range(sheet.ncols)] for r in range(sheet.nrows)]
        if not grid:
            continue
        header_index = pick_header_row(grid)
        tables.append(
            SheetTable(
                file_name=file_name,
                sheet_name=sheet.name,
                header=_dedupe_header(grid[header_index]),
                rows=grid[header_index + 1 :],
                header_row_index=header_index + 1,
            )
        )
    return tables


def read_csv(file_name: str, data: bytes) -> list[SheetTable]:
    text = None
    for encoding in ("utf-8-sig", "gb18030", "utf-16", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        raise RuntimeError("无法识别文本编码")
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    grid = [row for row in csv.reader(io.StringIO(text), dialect)]
    if not grid:
        return []
    header_index = pick_header_row(grid)
    return [
        SheetTable(
            file_name=file_name,
            sheet_name=os.path.splitext(os.path.basename(file_name))[0],
            header=_dedupe_header(grid[header_index]),
            rows=grid[header_index + 1 :],
            header_row_index=header_index + 1,
        )
    ]


def read_docx(file_name: str, data: bytes) -> tuple[list[SheetTable], TextDocument | None]:
    try:
        import docx  # type: ignore
    except ImportError as exc:
        raise RuntimeError("读取 .docx 需要安装 python-docx") from exc
    document = docx.Document(io.BytesIO(data))
    tables = []
    for index, table in enumerate(document.tables):
        grid = []
        for row in table.rows:
            seen: set[int] = set()
            values = []
            for cell in row.cells:
                if id(cell._tc) in seen:  # 合并单元格会重复出现同一对象
                    values.append(None)
                    continue
                seen.add(id(cell._tc))
                values.append(cell.text)
            grid.append(values)
        if not grid:
            continue
        header_index = pick_header_row(grid)
        tables.append(
            SheetTable(
                file_name=file_name,
                sheet_name=f"表{index + 1}",
                header=_dedupe_header(grid[header_index]),
                rows=grid[header_index + 1 :],
                header_row_index=header_index + 1,
            )
        )
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    return tables, TextDocument(file_name=file_name, kind="docx", text=text)


def read_pdf(file_name: str, data: bytes) -> tuple[list[SheetTable], TextDocument | None]:
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:
        raise RuntimeError("读取 PDF 需要安装 pdfplumber") from exc
    tables: list[SheetTable] = []
    chunks: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            chunks.append(page.extract_text() or "")
            for table_index, grid in enumerate(page.extract_tables() or [], start=1):
                if not grid:
                    continue
                header_index = pick_header_row(grid)
                tables.append(
                    SheetTable(
                        file_name=file_name,
                        sheet_name=f"第{page_index}页表{table_index}",
                        header=_dedupe_header(grid[header_index]),
                        rows=grid[header_index + 1 :],
                        header_row_index=header_index + 1,
                    )
                )
    text = "\n".join(chunks).strip()
    if not text and not tables:
        raise RuntimeError("PDF 无可提取文字，可能是扫描件，需要先 OCR")
    return tables, TextDocument(file_name=file_name, kind="pdf", text=text)
